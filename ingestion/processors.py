import json
import os
import re
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional

class BaseProcessor(ABC):
    @abstractmethod
    def can_handle(self, filename: str) -> bool:
        pass

    @abstractmethod
    def process(self, file_path: str) -> Dict[str, Any]:
        """Returns a dict with 'content' and 'metadata' and 'logs'."""
        pass

    def clean_text(self, text: str) -> str:
        # Basic cleaning: normalize whitespace, remove non-printable
        text = re.sub(r'\s+', ' ', text).strip()
        return text

class TextProcessor(BaseProcessor):
    def can_handle(self, filename: str) -> bool:
        return filename.lower().endswith(('.txt', '.md', '.py', '.log'))

    def process(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        cleaned_content = self.clean_text(content)
        return {
            "content": cleaned_content,
            "metadata": {"type": "text", "original_size": len(content)},
            "logs": ["Read as plain text", f"Normalized whitespace, reduced {len(content) - len(cleaned_content)} characters"]
        }

class JsonProcessor(BaseProcessor):
    def can_handle(self, filename: str) -> bool:
        return filename.lower().endswith('.json')

    def process(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        logs = ["Parsed JSON file"]
        
        # Heuristic for ChatGPT Export
        if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict) and 'title' in data[0] and 'mapping' in data[0]:
            return self._process_chatgpt(data, logs)
        
        # Heuristic for generic list of items (e.g., Instagram comments, Google Takeout)
        if isinstance(data, list) and len(data) > 0:
            return self._process_generic_list(data, logs)
            
        if isinstance(data, dict):
            # Check if it's a nested dict with a list of items
            for key, value in data.items():
                if isinstance(value, list) and len(value) > 0 and isinstance(value[0], (dict, str)):
                    logs.append(f"Found list of items in key: {key}")
                    return self._process_generic_list(value, logs)

        # Generic JSON
        content = json.dumps(data, indent=2)
        return {
            "content": content,
            "metadata": {"type": "json"},
            "logs": logs + ["Treated as generic JSON structure"]
        }

    def _process_generic_list(self, data: List[Any], logs: List[str]) -> Dict[str, Any]:
        logs.append(f"Processing generic list of {len(data)} items")
        text_lines = []
        for i, item in enumerate(data):
            if isinstance(item, str):
                text_lines.append(item)
            elif isinstance(item, dict):
                # Flatten the dict into a readable string
                parts = []
                for k, v in item.items():
                    if isinstance(v, (str, int, float, bool)):
                        parts.append(f"{k}: {v}")
                text_lines.append(" | ".join(parts))
        
        content = "\n".join(text_lines)
        return {
            "content": content,
            "metadata": {"type": "structured_list", "items": len(data)},
            "logs": logs + ["Extracted key-value pairs from list of objects"]
        }

    def _process_chatgpt(self, data: List[Dict], logs: List[str]) -> Dict[str, Any]:
        logs.append(f"Detected ChatGPT Export format ({len(data)} conversations)")
        full_text = []
        for conv in data:
            title = conv.get('title', 'Untitled')
            full_text.append(f"--- Conversation: {title} ---")
            mapping = conv.get('mapping', {})
            for node_id, node in mapping.items():
                message = node.get('message')
                if message and message.get('content') and message['content'].get('parts'):
                    role = message.get('author', {}).get('role', 'unknown')
                    parts = message['content']['parts']
                    text_parts = [p for p in parts if isinstance(p, str)]
                    if text_parts:
                        full_text.append(f"{role.upper()}: {' '.join(text_parts)}")
        
        content = "\n".join(full_text)
        return {
            "content": content,
            "metadata": {"type": "chatgpt_export", "conversations": len(data)},
            "logs": logs + ["Extracted text from conversation mapping"]
        }

class PdfProcessor(BaseProcessor):
    def can_handle(self, filename: str) -> bool:
        return filename.lower().endswith('.pdf')

    def process(self, file_path: str) -> Dict[str, Any]:
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            text = ""
            logs = [f"Detected PDF with {len(reader.pages)} pages"]
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {i+1} ---\n" + page_text
            
            return {
                "content": text,
                "metadata": {"type": "pdf", "pages": len(reader.pages)},
                "logs": logs + ["Extracted text using pypdf"]
            }
        except ImportError:
            return {
                "content": f"Error: pypdf not installed. Could not process {file_path}",
                "metadata": {"type": "pdf", "error": "missing_dependency"},
                "logs": ["FAILED: pypdf not installed"]
            }

class DocxProcessor(BaseProcessor):
    def can_handle(self, filename: str) -> bool:
        return filename.lower().endswith('.docx')

    def process(self, file_path: str) -> Dict[str, Any]:
        try:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return {
                "content": self.clean_text(text),
                "metadata": {"type": "docx", "paragraphs": len(doc.paragraphs)},
                "logs": [f"Detected DOCX with {len(doc.paragraphs)} paragraphs", "Extracted text using python-docx", "Normalized whitespace"]
            }
        except ImportError:
            return {
                "content": f"Error: python-docx not installed. Could not process {file_path}",
                "metadata": {"type": "docx", "error": "missing_dependency"},
                "logs": ["FAILED: python-docx not installed"]
            }

class HtmlProcessor(BaseProcessor):
    def can_handle(self, filename: str) -> bool:
        return filename.lower().endswith(('.html', '.htm'))

    def process(self, file_path: str) -> Dict[str, Any]:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
            
            # Remove scripts and styles
            for script in soup(["script", "style"]):
                script.extract()
            
            text = soup.get_text(separator='\n')
            cleaned_text = self.clean_text(text)
            
            return {
                "content": cleaned_text,
                "metadata": {"type": "html", "title": soup.title.string if soup.title else "No title"},
                "logs": ["Parsed HTML using BeautifulSoup", "Removed scripts and styles", "Normalized whitespace"]
            }
        except ImportError:
            return {
                "content": f"Error: beautifulsoup4 not installed. Could not process {file_path}",
                "metadata": {"type": "html", "error": "missing_dependency"},
                "logs": ["FAILED: beautifulsoup4 not installed"]
            }
